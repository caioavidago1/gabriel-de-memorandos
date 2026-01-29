from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import re


def export_memo_to_docx(memo_type: str, custom_fields: list, field_paragraphs: dict) -> BytesIO:
    """
    Exporta o memorando completo para formato DOCX.
    
    Args:
        memo_type: Tipo do memorando
        custom_fields: Lista de nomes das seções customizadas
        field_paragraphs: Dict com {field_name: [paragraphs]}
    
    Returns:
        BytesIO object com o documento DOCX
    """
    
    def add_formatted_paragraph(doc, text):
        """
        Adiciona parágrafo removendo markdown (** para negrito).
        
        Args:
            doc: Documento DOCX
            text: Texto com possível markdown
        
        Returns:
            Parágrafo do documento
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Remover ** do texto
        clean_text = text.replace('**', '')
        
        run = p.add_run(clean_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        
        return p
    
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== CAPA =====

    # Nome da empresa
    cover_spectra = doc.add_paragraph()
    run_spectra = cover_spectra.add_run("Spectra")
    run_spectra.font.size = Pt(32)
    run_spectra.font.bold = True
    run_spectra.font.name = 'Calibri'
    cover_spectra.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Espaço

    # Título do documento
    cover_title = doc.add_paragraph()
    run_title = cover_title.add_run(memo_type)
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_title.font.name = 'Calibri'
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Espaço

    # Mês e ano
    month_year = time.strftime('%B/%Y')
    cover_date = doc.add_paragraph()
    run_date = cover_date.add_run(month_year)
    run_date.font.size = Pt(12)
    run_date.font.italic = True
    run_date.font.name = 'Calibri'
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Quebra de página após a capa
    doc.add_page_break()

    # ===== FIM CAPA =====

    # Título do documento (repetido no início do conteúdo)
    title_para = doc.add_paragraph()
    run_title2 = title_para.add_run(memo_type)
    run_title2.font.size = Pt(12)
    run_title2.font.bold = True
    run_title2.font.name = 'Calibri'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Espaço
    
    # Adicionar cada seção customizada
    for field_name in custom_fields:
        # Título da seção
        heading_para = doc.add_paragraph()
        run_heading = heading_para.add_run(field_name)
        run_heading.font.size = Pt(12)
        run_heading.font.bold = True
        run_heading.font.name = 'Calibri'
        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Parágrafos da seção
        paragraphs = field_paragraphs.get(field_name, [])

        if paragraphs:
            for para_text in paragraphs:
                if para_text and para_text.strip():
                    # Usar função que processa markdown
                    add_formatted_paragraph(doc, para_text)
        else:
            # Se não tem conteúdo, adicionar placeholder
            p = doc.add_paragraph("[Conteúdo a ser adicionado]")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Texto justificado
            p.runs[0].italic = True
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(12)

        # Espaço entre seções
        doc.add_paragraph()
    
    # Se não tem seções, adicionar mensagem
    if not custom_fields:
        doc.add_paragraph("Nenhuma seção foi criada ainda. Adicione seções para gerar o memorando.")
    
    # Salvar em BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file
